"""
Phase 1: Extract Chapter 3 (Methodology) from the thesis docx to markdown.

Uses python-docx to locate the Chapter 3 heading and extract everything
until the next chapter-level heading. Preserves tables in pipe format
and notes figure captions.
"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


DOCX_PATH = Path(__file__).parents[2] / "Beley-Reyes_Thesis2-ACM.docx"
OUTPUT_PATH = Path(__file__).parents[1] / "outputs" / "chapter3_methodology.md"

# Heading styles that mark chapter boundaries (Heading 1 equivalents in ACM template)
CHAPTER_HEADING_STYLES = {
    "Heading 1", "heading 1",
    "ACM-Heading", "acm-heading",
}

# Patterns that indicate a Chapter 3 start
CH3_PATTERNS = [
    re.compile(r"^3\s*[\.\s]", re.IGNORECASE),
    re.compile(r"methodology", re.IGNORECASE),
    re.compile(r"chapter\s+3", re.IGNORECASE),
    re.compile(r"^research\s+design", re.IGNORECASE),
    re.compile(r"^3\s+methodology", re.IGNORECASE),
]

# Pattern that marks the start of Chapter 4 (end of Chapter 3)
CH4_PATTERNS = [
    re.compile(r"^4[\.\s]", re.IGNORECASE),
    re.compile(r"results", re.IGNORECASE),
    re.compile(r"chapter\s+4", re.IGNORECASE),
    re.compile(r"^discussion", re.IGNORECASE),
    re.compile(r"^evaluation", re.IGNORECASE),
]

HEADING_STYLE_MAP = {
    1: "#",
    2: "##",
    3: "###",
    4: "####",
    5: "#####",
}


def _heading_level(para) -> int | None:
    style_name = para.style.name if para.style else ""
    for i in range(1, 6):
        if f"Heading {i}" in style_name or f"heading {i}" in style_name:
            return i
    return None


def _is_chapter_boundary(para, patterns: list) -> bool:
    text = para.text.strip()
    lvl = _heading_level(para)
    if lvl not in (None, 1, 2):
        return False
    if lvl == 1:
        return any(p.search(text) for p in patterns)
    # Some ACM templates use style 2 for chapter titles
    if lvl == 2 and len(text) < 80:
        return any(p.search(text) for p in patterns)
    return False


def _table_to_md(table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


def _run_text(para) -> str:
    """Get full text of a paragraph, handling footnote refs gracefully."""
    parts = []
    for run in para.runs:
        parts.append(run.text)
    return "".join(parts)


def _para_to_md(para) -> str:
    text = _run_text(para).strip()
    if not text:
        return ""
    lvl = _heading_level(para)
    if lvl:
        prefix = HEADING_STYLE_MAP.get(lvl, "##")
        return f"{prefix} {text}"
    style = para.style.name if para.style else ""
    if "caption" in style.lower() or text.startswith("Figure") or text.startswith("Table"):
        return f"*{text}*"
    if "list" in style.lower() or "bullet" in style.lower():
        return f"- {text}"
    return text


class Chapter3Extractor:
    """Extract Chapter 3 content from the thesis docx as markdown."""

    def __init__(self, docx_path: Path):
        self.doc = Document(str(docx_path))
        self.paragraphs = self.doc.paragraphs
        self._table_index = self._build_table_index()

    def _build_table_index(self) -> dict[int, object]:
        """Map paragraph XML element position to tables that follow them."""
        body = self.doc.element.body
        table_map = {}
        children = list(body)
        for i, child in enumerate(children):
            if child.tag == qn("w:tbl"):
                # Find the preceding paragraph index
                table_map[i] = child
        return table_map

    def _iter_body_elements(self):
        """Yield (index, kind, object) for each body element."""
        body = self.doc.element.body
        para_iter = iter(self.doc.paragraphs)
        for i, child in enumerate(body):
            if child.tag == qn("w:p"):
                try:
                    para = next(para_iter)
                    yield (i, "para", para)
                except StopIteration:
                    break
            elif child.tag == qn("w:tbl"):
                # Find the table object
                for tbl in self.doc.tables:
                    if tbl._tbl is child:
                        yield (i, "table", tbl)
                        break

    def extract(self) -> str:
        in_ch3 = False
        lines = []

        for idx, kind, obj in self._iter_body_elements():
            if kind == "para":
                para = obj
                text = para.text.strip()

                if not in_ch3:
                    if _is_chapter_boundary(para, CH3_PATTERNS):
                        in_ch3 = True
                        md = _para_to_md(para)
                        if md:
                            lines.append(md)
                    continue

                # Check if we've hit the next chapter
                if _is_chapter_boundary(para, CH4_PATTERNS):
                    break

                md = _para_to_md(para)
                if md:
                    lines.append(md)
                elif not text:
                    lines.append("")  # Preserve blank lines

            elif kind == "table" and in_ch3:
                lines.append("")
                lines.append(_table_to_md(obj))
                lines.append("")

        return "\n".join(lines)


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    extractor = Chapter3Extractor(DOCX_PATH)
    content = extractor.extract()
    if not content.strip():
        raise RuntimeError(
            "No Chapter 3 content found. The heading detection may need adjustment."
        )
    OUTPUT_PATH.write_text(content, encoding="utf-8")
    lines = content.splitlines()
    print(f"Extracted {len(lines)} lines to {OUTPUT_PATH}")
    print(f"Preview (first 10 lines):\n" + "\n".join(lines[:10]))


if __name__ == "__main__":
    main()
