"""Shared helpers for programmatic revision of the thesis docx (python-docx)."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

DOC_PATH = r"D:\Developer\Projects\THESIS-Utilizing-ML-to-Solve-the-IPPP\Beley-Reyes_Thesis2-ACM.docx"

BODY = "Thesis: Paragraph Text"
CAPTION = "Caption"


def load():
    return Document(DOC_PATH)


def save(doc):
    doc.save(DOC_PATH)


def iter_blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def find_par(doc, prefix, style=None, nth=0):
    """Find the nth paragraph whose text starts with prefix (whitespace-normalized)."""
    hits = 0
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            if " ".join(block.text.split()).startswith(" ".join(prefix.split())):
                if style and (not block.style or block.style.name != style):
                    continue
                if hits == nth:
                    return block
                hits += 1
    raise LookupError(f"paragraph not found: {prefix!r}")


def par_exists(doc, prefix):
    try:
        find_par(doc, prefix)
        return True
    except LookupError:
        return False


def find_table(doc, header_prefix, nth=0):
    """Find the nth table whose first-row text starts with header_prefix."""
    hits = 0
    for block in iter_blocks(doc):
        if isinstance(block, Table):
            first = " | ".join(c.text.strip() for c in block.rows[0].cells)
            if first.startswith(header_prefix):
                if hits == nth:
                    return block
                hits += 1
    raise LookupError(f"table not found: {header_prefix!r}")


def _new_par(doc, text="", style=BODY):
    p = OxmlElement("w:p")
    par = Paragraph(p, doc)
    if style:
        par.style = doc.styles[style]
    if text:
        par.add_run(text)
    return par


def insert_par_after(doc, anchor_element, text="", style=BODY):
    """Insert a new paragraph after anchor_element (a w:p or w:tbl element).
    Returns the new Paragraph."""
    par = _new_par(doc, text, style)
    anchor_element.addnext(par._p)
    return par


def insert_par_before(doc, anchor_element, text="", style=BODY):
    par = _new_par(doc, text, style)
    anchor_element.addprevious(par._p)
    return par


def set_par_text(par, text, style=None):
    """Replace all content of a paragraph with a single plain run (drops fields)."""
    for child in list(par._p):
        if child.tag != qn("w:pPr"):
            par._p.remove(child)
    if style:
        par.style = par.part.document.styles[style] if hasattr(par, "part") else style
    par.add_run(text)
    return par


def delete_par(par):
    par._p.getparent().remove(par._p)


def insert_table_after(doc, anchor_element, rows, style="Grid Table Light",
                       bold_header=True):
    """Create a table from rows (list of lists of str) and place it after
    anchor_element. Returns the Table."""
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = doc.styles[style]
    tbl.autofit = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = tbl.cell(r, c)
            cell.text = str(val)
            if bold_header and r == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    anchor_element.addnext(tbl._tbl)
    return tbl


def page_break_before(par, enable=True):
    """Set w:pageBreakBefore on a paragraph."""
    pPr = par._p.get_or_add_pPr()
    for el in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(el)
    if enable:
        el = OxmlElement("w:pageBreakBefore")
        pPr.insert(0, el)


def append_text(par, text):
    """Append a run of text to an existing paragraph (joined with a space)."""
    if par.text and not par.text.endswith((" ", " ")):
        text = " " + text
    par.add_run(text)


def move_before(element, anchor_element):
    anchor_element.addprevious(element)


def move_after(element, anchor_element):
    anchor_element.addnext(element)
